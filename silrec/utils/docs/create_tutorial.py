from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

HEADER_BG = '2C3E50'
WHITE = 'FFFFFF'
BLACK = '000000'
GRAY = '7F8C8D'
LIGHT_BG = 'F0F2F5'
TIP_BG = 'EBF5FB'
NOTE_BG = 'FEF9E7'

def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(HEADER_BG)

def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(HEADER_BG)

def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string('34495E')

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def step(number, text):
    p = doc.add_paragraph()
    run = p.add_run(f'{number}. ')
    run.font.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p

def bullet(text, indent=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if indent:
        p.paragraph_format.left_indent = Inches(0.5 + indent * 0.3)
    p.paragraph_format.space_after = Pt(2)
    return p

def screenshot(filename, description):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'[Screenshot: {filename}]')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x5D, 0xAD, 0xEC)
    run2 = p.add_run(f' \u2014 {description}')
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

def tip_box(text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TIP_BG}" w:val="clear"/>')
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'\U0001f4a1 Tip: ')
    run.font.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)

def note_box(text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{NOTE_BG}" w:val="clear"/>')
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'\u2139 Note: ')
    run.font.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for ci, h in enumerate(headers):
        cell = table.cell(0, ci)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEADER_BG}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        for ci, text in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9)
            if ri % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA"/>')
                cell._tc.get_or_add_tcPr().append(shading)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if col_widths:
        for ci, w in enumerate(col_widths):
            table.columns[ci].width = Emu(int(w * 914400))
    doc.add_paragraph()
    return table

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="BDC3C7"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

# ============================================================
# TITLE PAGE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
run = p.add_run('SILREC\nSelf-Guided Tutorial')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string(HEADER_BG)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Silvicultural Recording System\nLearn the application step by step')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor.from_string(GRAY)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(60)
run3 = p3.add_run('Follow this guide to explore each feature of SILREC.\nEach section includes step-by-step instructions and placeholders for screenshots.')
run3.font.size = Pt(12)
run3.font.color.rgb = RGBColor.from_string('5D6D7E')

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
h1('Table of Contents')
toc_items = [
    '1. Getting Started \u2014 Logging In and Overview',
    '2. The Workflow \u2014 Understanding Proposals',
    '3. New Application \u2014 Uploading and Processing a Shapefile',
    '4. Entering Data \u2014 Cohorts, Treatments, and More',
    '5. The Treatments Page \u2014 Standalone Data Entry',
    '6. The Map \u2014 Spatial View of Polygons',
    '7. Reports \u2014 Generating and Downloading',
    '8. Search \u2014 Finding Records',
    '9. System Maintenance \u2014 Admin Tasks',
    '10. Quick Reference \u2014 URLs and Navigation',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(6)
    run = p.runs[0]
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(HEADER_BG)

doc.add_page_break()

# ============================================================
# 1. GETTING STARTED
# ============================================================
h1('1. Getting Started \u2014 Logging In and Overview')

h2('1.1 Logging In')
step(1, 'Open your browser and navigate to the SILREC URL provided by your administrator.')
step(2, 'Click the Log in button. If using SSO (Single Sign-On), you will be redirected to the DBCA authentication page.')
step(3, 'Enter your DBCA credentials and log in.')
step(4, 'After login, you arrive at the Proposals dashboard \u2014 the main landing page.')
screenshot('login-page.png', 'SILREC login screen')

h2('1.2 Understanding the Navigation Menu')
body('The top navigation bar gives you access to all major sections:')
bullet('Home / Proposals \u2014 the main dashboard')
bullet('Treatments \u2014 standalone treatments datatable')
bullet('Map \u2014 spatial polygon viewer')
bullet('Reports \u2014 report generator')
bullet('Search \u2014 text, user, and reference search')
bullet('Options \u2014 generated reports, system maintenance (if permitted)')
note_box('Your menu items may differ depending on your assigned role. If you don\'t see a menu item, your role may not have permission for that section.')

h2('1.3 Your Role Determines What You Can Do')
body('SILREC has four user groups. Your group membership controls what you can see and do:')
add_table(
    ['Role', 'What You Can Do'],
    [
        ['User', 'View proposals, treatments, map, reports, search. Cannot create or edit anything.'],
        ['Operator', 'Create proposals, upload shapefiles, enter cohort/treatment data, send to reviewer. Core data entry role.'],
        ['Reviewer', 'Review proposal data, approve or send back to operator.'],
        ['Silrec Admin', 'Full access \u2014 manage users, configure system, bypass restrictions.'],
    ],
    col_widths=[1.5, 5.0]
)
tip_box('Hover over any button or menu item. If it is grayed out, your current role does not permit that action for the current proposal status.')

screenshot('navigation-menu.png', 'Top navigation bar showing available sections')

divider()

# ============================================================
# 2. THE WORKFLOW
# ============================================================
h1('2. The Workflow \u2014 Understanding Proposals')

h2('2.1 What is a Proposal?')
body('A proposal is the container for a shapefile upload and all related data. Every shapefile you upload becomes part of a proposal, and the proposal moves through a standard lifecycle.')

h2('2.2 The Processing Status Lifecycle')
body('Every proposal moves through the following statuses in order:')

add_table(
    ['Status', 'What Happens'],
    [
        ['Draft', 'Proposal is created. No shapefile has been uploaded yet.'],
        ['Processing Shapefile', 'Shapefile is being processed by the system. No other proposal can be processed at the same time.'],
        ['With Operator', 'The Operator enters cohort data, treatments, prescriptions, and survey documents.'],
        ['With Reviewer', 'The Reviewer checks the data, approves it, or sends it back for changes.'],
        ['Review Completed', 'The Reviewer has signed off. The proposal is complete.'],
        ['Amendment Required', 'Changes were requested \u2014 the proposal goes back to the Operator.'],
        ['Declined / Discarded', 'The proposal was rejected or abandoned.'],
    ],
    col_widths=[1.8, 4.7]
)

body('Flow diagram:')
code('Draft')
code('  \u2502  (upload & process shapefile)')
code('  \u250c')
code('Processing Shapefile')
code('  \u2502  (processing completes)')
code('  \u250c')
code('With Operator')
code('  \u2502  (operator completes data entry)')
code('  \u250c')
code('With Reviewer')
code('  \u2502  (reviewer approves/rejects)')
code('  \u250c')
code('Review Completed')

screenshot('workflow-diagram.png', 'Visual workflow diagram showing status transitions')

h2('2.3 Who Can Move Between Statuses')
body('Each transition from one status to another requires a specific role:')

add_table(
    ['From \u2192 To', 'Transition', 'Who Can Do It'],
    [
        ['Draft \u2192 Processing Shapefile', 'Upload & process shapefile', 'Operator'],
        ['Processing Shapefile \u2192 With Operator', 'Processing complete', 'Operator, Silrec Admin'],
        ['With Operator \u2192 With Reviewer', 'Send to reviewer', 'Operator, Silrec Admin'],
        ['With Operator \u2192 Draft', 'Return to draft', 'Operator, Silrec Admin'],
        ['With Reviewer \u2192 Review Completed', 'Approve', 'Reviewer, Silrec Admin'],
        ['With Reviewer \u2192 With Operator', 'Send back for changes', 'Reviewer, Silrec Admin'],
        ['Review Completed \u2192 With Reviewer', 'Reopen', 'Silrec Admin'],
    ],
    col_widths=[1.8, 2.0, 2.0]
)

screenshot('workflow-buttons.png', 'Workflow action buttons at the bottom of a proposal detail page')

h2('2.4 How to Change a Proposal\'s Status')
step(1, 'Open the proposal by clicking its Lodgement Number on the Proposals dashboard.')
step(2, 'Scroll to the bottom of the proposal detail page. You will see a set of action buttons.')
step(3, 'Click the button that matches the desired transition (e.g., Send to Review).')
step(4, 'If prompted, enter a comment explaining the transition.')
step(5, 'The status updates immediately. A version comment is logged automatically.')
tip_box('Only buttons for permitted transitions appear. If you don\'t see a button, your role may not be allowed to perform that transition at the current status.')

divider()

# ============================================================
# 3. NEW APPLICATION
# ============================================================
h1('3. New Application \u2014 Uploading and Processing a Shapefile')

h2('3.1 Creating a New Proposal')
step(1, 'On the Proposals dashboard, click the New Proposal button.')
step(2, 'Select the Application Type from the dropdown.')
step(3, 'The proposal is created in Draft status and appears in the datatable.')
screenshot('new-proposal-button.png', 'New Proposal button on the dashboard')

h2('3.2 Uploading a Shapefile')
step(1, 'Click the proposal\'s Lodgement Number to open the detail page.')
step(2, 'Find the Shapefile section.')
step(3, 'Click Upload Shapefile and select your .zip file containing the shapefile components (.shp, .dbf, .prj, .shx).')
step(4, 'You can upload a single .zip or individual component files.')
step(5, 'After upload, the shapefile attributes are displayed for validation.')
note_box('Accepted formats: .zip containing all shapefile components, or individual .shp / .dbf / .prj / .shx files.')
screenshot('shapefile-upload.png', 'Shapefile upload section on the proposal detail page')

h2('3.3 Shapefile Attribute Validation')
body('When the shapefile is uploaded, the system automatically validates its attributes:')
bullet('Mandatory fields \u2014 must be present; missing fields are flagged with an error.')
bullet('Data types \u2014 each attribute is checked against the expected type (text, number, date, etc.).')
bullet('Reserved fields \u2014 certain field names cannot be used because the system needs them.')
body('If validation fails, fix the shapefile and re-upload. If it passes, you can proceed to processing.')
screenshot('shapefile-validation.png', 'Shapefile attribute validation results')

h2('3.4 Processing the Shapefile')
step(1, 'Once attribute validation passes, click Process Shapefile.')
step(2, 'The system runs the ShapefileSliversMerger, which:')
bullet('Projects the geometry to the working coordinate system')
bullet('Detects and dissolves sliver polygons (very thin polygons below the area/length threshold)')
bullet('Detects and resolves polygon overlaps')
bullet('Writes the processed data into the polygon, cohort, and assign_cht_to_ply tables')
step(3, 'Processing is tracked via Processing Runs and Savepoint Records, showing each iteration.')
step(4, 'While processing, the proposal is locked \u2014 no other proposal can be processed at the same time (other proposals show a lock icon).')
screenshot('process-shapefile-button.png', 'Process Shapefile button')
screenshot('processing-lock-icon.png', 'Lock icon on other proposals during processing')

h2('3.5 Keep or Revert After Processing')
body('Once processing completes, you have two options:')
bullet('Keep \u2014 confirms the processed data is correct and transitions the proposal to With Operator status.')
bullet('Revert \u2014 undoes all changes and restores the polygon, cohort, and assign_cht_to_ply tables to their pre-processing state.')
body('The revert mechanism uses savepoints by default (backup copies of the three core tables). If configured, it can also use pg_dump.')
screenshot('keep-revert-buttons.png', 'Keep and Revert buttons after processing')

h2('3.6 The Proposals Dashboard (Datatable)')
body('The main Proposals dashboard lists all proposals in an interactive table:')
bullet('Filters \u2014 collapse/expand the filter panel above the table: Type (application type), Status (processing status), Lodged From / Lodged To (date range).')
bullet('Sorting \u2014 click any column header to sort ascending or descending.')
bullet('Search \u2014 the search box at the top searches across all visible columns.')
bullet('Lock indicators \u2014 proposals currently being processed appear first with a lock icon.')
screenshot('proposals-datatable.png', 'Proposals datatable with filters and search')

h2('3.7 Navigating a Proposal Record')
step(1, 'Click any proposal\'s Lodgement Number to open its detail page.')
step(2, 'The detail page shows:')
bullet('Proposal header with lodgement number and current processing status')
bullet('Status transition comment alerts (shown if the status was recently reverted)')
bullet('An Application Form with tabs for different data sections')
bullet('A fixed bottom navbar with workflow action buttons')
screenshot('proposal-detail.png', 'Proposal detail page showing header, tabs, and action buttons')

divider()

# ============================================================
# 4. ENTERING DATA
# ============================================================
h1('4. Entering Data \u2014 Cohorts, Treatments, and More')

h2('4.1 Cohorts')
body('Each polygon can have one or more cohorts (species groups). A cohort represents a group of trees with similar characteristics within a polygon.')
step(1, 'Navigate to the Cohorts section within the proposal detail page.')
step(2, 'Click Add Cohort or the edit icon next to an existing cohort.')
step(3, 'Fill in the cohort details: species, age class, structure, density, etc.')
step(4, 'Click Save.')
screenshot('cohort-form.png', 'Cohort edit form with species and age class fields')

h2('4.2 Treatments')
body('Each cohort can have one or more treatments (silvicultural activities). A treatment records what action was taken, when, and by whom.')
step(1, 'Navigate to the Treatments section within a cohort or proposal.')
step(2, 'Click Add Treatment or edit an existing one.')
step(3, 'Fill in the treatment details:')
bullet('Task \u2014 the type of silvicultural activity')
bullet('Classification \u2014 parent category for the task')
bullet('Status \u2014 planned, in progress, completed, etc.')
bullet('Planned Year / Month \u2014 when the treatment is scheduled')
bullet('Completed Date \u2014 when it was actually done')
bullet('Machine \u2014 equipment used')
bullet('Operator \u2014 person who performed the work')
bullet('Area \u2014 hectares treated')
bullet('Cost \u2014 cost information')
step(4, 'Click Save.')
screenshot('treatment-form.png', 'Treatment edit form with task, status, dates, and machine fields')

h2('4.3 Prescriptions')
body('Prescriptions are silvicultural instructions linked to a treatment.')
step(1, 'Within a treatment detail page, find the Prescriptions section.')
step(2, 'Click Add Prescription.')
step(3, 'Enter the prescription details and save.')
screenshot('prescription-form.png', 'Prescription form linked to a treatment')

h2('4.4 Silviculturist Comments')
step(1, 'Within a treatment detail page, find the Silviculturist Comments section.')
step(2, 'Add free-text notes about the treatment.')
step(3, 'Save your comments.')
screenshot('silviculturist-comments.png', 'Silviculturist comments section on treatment detail')

h2('4.5 Survey Assessment Documents')
step(1, 'Within a treatment detail page, find the Survey Assessment Documents section.')
step(2, 'Upload documents such as shapefiles, PDFs, or images.')
step(3, 'The uploaded files are linked to the treatment record.')
screenshot('survey-documents-upload.png', 'Survey assessment document upload section')

h2('4.6 Form Validation Rules')
body('When you save a form, the system checks Form Validation Rules configured by the Silrec Admin:')
bullet('Some fields are required only at certain statuses (e.g., treatment cost may be mandatory before sending to review).')
bullet('If a required field is empty, the form will show a validation error and will not save.')
bullet('The Silrec Admin configures these rules at /admin/ under Silrec > Form validation rules.')
screenshot('admin-form-validation-rules.png', 'Form Validation Rules admin configuration')

divider()

# ============================================================
# 5. TREATMENTS PAGE
# ============================================================
h1('5. The Treatments Page \u2014 Standalone Data Entry')

h2('5.1 Accessing the Treatments Page')
body('Click Treatments in the top navigation menu. This opens a standalone treatments datatable that shows all treatments across all proposals, independent of any specific proposal.')

h2('5.2 Filtering Treatments')
body('The filter panel (collapsible) lets you narrow down treatments:')
bullet('Task Classification \u2014 selecting a parent classification filters the Task dropdown')
bullet('Task \u2014 searchable dropdown that updates based on the selected classification')
bullet('Status \u2014 filter by treatment status (planned, completed, etc.)')
bullet('Plan Year / Plan Month \u2014 filter by scheduled date')
bullet('Complete Date Range \u2014 filter by actual completion date')
bullet('Machine \u2014 filter by equipment used')
bullet('Operator \u2014 filter by assigned person')
bullet('Refresh button \u2014 reload the data')
screenshot('treatments-datatable.png', 'Treatments datatable with filter panel expanded')

h2('5.3 Editing a Treatment')
step(1, 'Click a treatment record to open the Treatment Detail page (/internal/treatment/{id}/).')
step(2, 'Here you can:')
bullet('Edit treatment fields (task, status, dates, area, cost, etc.)')
bullet('Add Treatment Extras \u2014 additional line items for a treatment')
bullet('Add Prescriptions \u2014 silvicultural instructions')
bullet('Add Silviculturist Comments')
bullet('Upload Survey Assessment Documents')
step(3, 'Each sub-form has its own Save button.')
step(4, 'Form validation rules (configured by the admin) apply here too.')
screenshot('treatment-detail.png', 'Treatment detail page with all sub-sections')
screenshot('treatment-extras-form.png', 'Treatment extras form for additional line items')

divider()

# ============================================================
# 6. THE MAP
# ============================================================
h1('6. The Map \u2014 Spatial View of Polygons')

h2('6.1 Opening the Map')
body('Click Map in the top navigation menu to open the spatial view at /internal/map.')

h2('6.2 Navigating the Map')
body('The map displays all polygons as a spatial layer. You can:')
bullet('Pan and zoom using mouse controls or the map toolbar.')
bullet('Search by text \u2014 find polygons by compartment, block, district, FEA ID, or objective code.')
bullet('Use the filter panel to narrow down polygons by:')
bullet('Compartment', indent=1)
bullet('Block', indent=1)
bullet('District', indent=1)
bullet('Objective classification', indent=1)
bullet('Treatment status', indent=1)
bullet('Created date range', indent=1)
bullet('Post-2024 filter \u2014 automatically filters to the configured plan period.')
screenshot('map-view.png', 'Map view with polygon layer and filter panel')

h2('6.3 Interacting with Features')
step(1, 'Click any polygon on the map to open a feature popup.')
step(2, 'The popup shows:')
bullet('Block and compartment number')
bullet('FEA ID')
bullet('Area in hectares')
bullet('Objective code')
bullet('A link to the polygon detail page')
step(3, 'Click the link to navigate to the Polygon Detail page, where you can view attributes and access linked cohorts and treatments.')
screenshot('map-feature-popup.png', 'Feature popup on the map with polygon details')

h2('6.4 Merge and Cut Tools')
body('If enabled by the administrator, the map also provides:')
bullet('Merge Polygon Tool \u2014 select two adjacent polygons and merge them into one.')
bullet('Cut Polygon Tool \u2014 draw a cutting line to split a polygon into two.')
note_box('These tools are controlled by the SHOW_MERGE_POLYGON_TOOL and SHOW_CUT_POLYGON_TOOL settings. If you don\'t see them, they may not be enabled for your instance.')
screenshot('merge-polygon-tool.png', 'Merge polygon tool in action')
screenshot('cut-polygon-tool.png', 'Cut polygon tool with cutting line')

divider()

# ============================================================
# 7. REPORTS
# ============================================================
h1('7. Reports \u2014 Generating and Downloading')

h2('7.1 Accessing the Reports Page')
body('Click Reports in the top navigation menu to open /internal/reports.')

h2('7.2 Step-by-Step Report Generation')
h3('Step 1: Select Report')
step(1, 'Choose a report type from the dropdown.')
step(2, 'Each report has a description explaining what data it shows.')
screenshot('report-step1-select.png', 'Report generator Step 1: selecting a report type')

h3('Step 2: Set Parameters')
body('Depending on the report, you may see different parameter types:')
bullet('Single Select \u2014 choose one option from a dropdown.')
bullet('Multi-Select \u2014 choose multiple options (Select2 tag mode).')
bullet('Year Enhanced \u2014 enter a single year (2024), comma-separated years (2024,2023), or a range (2020-2024).')
bullet('Text / Number \u2014 free-text or numeric input.')
bullet('Date Picker \u2014 select a date.')
bullet('Month Select \u2014 pick a month.')
bullet('Range \u2014 two number inputs for BETWEEN queries.')
bullet('Custom WHERE \u2014 advanced users can add field/operator/value filters.')
screenshot('report-step2-parameters.png', 'Report generator Step 2: setting query parameters')

h3('Step 3: Select Export Format')
step(1, 'Choose your output format:')
bullet('Excel (.xlsx)')
bullet('CSV (.csv)')
bullet('PDF (.pdf)')
bullet('Shapefile (.shp) \u2014 downloads as a zipped shapefile')
step(2, 'Click Generate Report to download, or click Preview Data first to see results in a modal.')
screenshot('report-step3-format.png', 'Report generator Step 3: selecting export format')

h2('7.3 Previewing Data Before Generating')
step(1, 'Click Preview Data on the report generator page.')
step(2, 'A modal opens showing:')
bullet('The generated SQL query (in a collapsible section)')
bullet('The first rows of data')
step(3, 'This lets you verify that your parameters produce the expected results before generating the full report.')
screenshot('report-preview-modal.png', 'Preview Data modal with SQL and sample results')

h2('7.4 Downloading and Managing Reports')
h3('Immediate Download')
body('After clicking Generate Report, the report is generated asynchronously. Once ready, it downloads directly to your browser. The filename includes the report name and a timestamp.')

h3('Re-downloading Past Reports')
step(1, 'Go to Options > Generated Reports (or navigate to /mgt-commands/generated-reports/).')
step(2, 'A list shows all previously generated report files with timestamps.')
step(3, 'Click a filename to download it again.')
step(4, 'Reports are kept for a configurable number of days (default: 20).')
step(5, 'Silrec Admin users can delete individual report files if needed.')
screenshot('generated-reports-list.png', 'Generated Reports list page')

h2('7.5 Admin Configuration (for Reference)')
body('Report definitions are configured by the Silrec Admin at /admin/ under Proposals > Sql reports. Each report defines:')
bullet('A SQL query to fetch data')
bullet('Accepts parameters (WHERE clauses) defined as JSON')
bullet('Allowed export formats')
bullet('Column display settings')
bullet('Allowed user groups')
body('Admins can clone an existing report to create a new one with the same structure.')
screenshot('admin-sql-reports.png', 'SQL Reports admin list')
screenshot('admin-sql-report-edit.png', 'SQL Report admin edit form')

h2('7.6 PDF Word Templates')
body('PDF reports use Word (.docx) templates:')
bullet('Templates are uploaded and versioned in the admin interface.')
bullet('One template per report is marked as current (the active version).')
bullet('When generating a PDF, the system renders the current template with the report data.')

divider()

# ============================================================
# 8. SEARCH
# ============================================================
h1('8. Search \u2014 Finding Records')

h2('8.1 Accessing Search')
body('Click Search in the top navigation menu to open /internal/search. Three search modes are available via tabs.')

h2('8.2 Text Search')
step(1, 'Enter a search string (minimum 2 characters).')
step(2, 'Select a model type to search within: polygons, cohorts, treatments, proposals, etc. Each option shows how many search fields it covers.')
step(3, 'Choose a match type:')
bullet('Contains (default) \u2014 text appears anywhere in the field')
bullet('Exact Match \u2014 must match exactly')
bullet('Starts With / Ends With')
step(4, 'Optionally set a Date From / Date To range.')
step(5, 'Click Search to display results in a table.')
step(6, 'Results show matching records with links to their detail pages.')
screenshot('search-text.png', 'Text Search page with parameters')
screenshot('search-text-results.png', 'Text Search results table')

h2('8.3 User Search')
body('Search for users by name, email, or username. Useful for finding operator or reviewer assignments.')
screenshot('search-user.png', 'User Search page')

h2('8.4 Reference Number Search')
body('Quick lookup by lodgement number or reference number.')
screenshot('search-reference.png', 'Reference Number Search page')

h2('8.5 Admin Configuration (for Reference)')
body('The Silrec Admin configures searchable models and fields at /admin/ under Proposals > Text search model configs:')
bullet('Each config entry defines which model to search, which fields to search, and how results are displayed.')
bullet('Search field display names can be customised via Text search field displays.')
note_box('Text search model configs use app_label forest_blocks.<Model> (not silrec). Proposals use app_label silrec.Proposal. If a search returns no results, check that the model config is set up correctly.')

divider()

# ============================================================
# 9. SYSTEM MAINTENANCE
# ============================================================
h1('9. System Maintenance \u2014 Admin Tasks')
note_box('This section is only accessible to superusers and Silrec Admin users.')

h2('9.1 Accessing System Maintenance')
body('Navigate to /mgt-commands/ or go to Options > System Maintenance from the top menu.')
screenshot('system-maintenance.png', 'System Maintenance page')

h2('9.2 What You Can Do')
body('The maintenance page displays:')
bullet('DB Dumps section \u2014 list of database backup files (pg_dump outputs) available for download. Admins can delete individual files.')
bullet('Generated Reports section \u2014 previously generated reports available for re-download.')
bullet('Cleanup tasks \u2014 run automatically on a schedule or manually via management commands.')

h2('9.3 Maintenance Tasks')
add_table(
    ['Task', 'Command', 'Schedule'],
    [
        ['Database backup', 'python manage.py db_dump', 'Cron (configured via CRON_RUN_AT_TIMES)'],
        ['Clean old reports', 'python manage.py cleanup_generated_reports --keep-days 20', 'Cron'],
        ['Manual revert (pg_dump)', 'python manage.py manual_pgrestore_revert --proposal-id=X', 'As needed'],
        ['Restore shapefile dump', 'python manage.py restore_shapefile_dump --proposal-id=X', 'As needed'],
    ],
    col_widths=[1.8, 3.0, 2.0]
)
screenshot('db-dumps-list.png', 'DB Dumps list page')

h2('9.4 Configuration Options')
body('The following can be set via environment variables:')
bullet('DB_DUMPS_DIR \u2014 where database dumps are stored (default: file_exports/db_dumps)')
bullet('REPORT_EXPORT_DIR \u2014 where generated reports are stored (default: file_exports/generated_reports)')
bullet('REPORT_RETENTION_DAYS \u2014 how long to keep generated reports (default: 20)')
bullet('SHAPEFILE_EXPORT_KEEP \u2014 number of shapefile exports to retain (default: 10)')

divider()

# ============================================================
# 10. QUICK REFERENCE
# ============================================================
h1('10. Quick Reference \u2014 URLs and Navigation')

h2('Navigation Paths')
add_table(
    ['Page', 'URL', 'Description'],
    [
        ['Home / Proposals', '/ or /internal/', 'Proposals datatable + Map tab'],
        ['Proposal Detail', '/internal/proposal/{id}/', 'View and edit a proposal'],
        ['Treatments', '/internal/treatments/', 'Standalone treatments datatable'],
        ['Treatment Detail', '/internal/treatment/{id}/', 'View and edit a treatment record'],
        ['Map', '/internal/map/', 'Spatial polygon viewer'],
        ['Reports', '/internal/reports/', 'Report generator'],
        ['Search', '/internal/search/', 'Text, user, and reference search'],
        ['System Maintenance', '/mgt-commands/', 'DB dumps and generated reports management'],
        ['Django Admin', '/admin/', 'Full system configuration (Silrec Admin only)'],
        ['Generated Reports', '/mgt-commands/generated-reports/', 'Re-download past reports'],
        ['DB Dumps', '/mgt-commands/db-dumps/', 'Download database backup files'],
    ],
    col_widths=[1.5, 2.5, 3.5]
)

h2('User Group Configuration')
body('To add users to groups (requires Silrec Admin access):')
step(1, 'Log in to /admin/.')
step(2, 'Go to Authentication and Authorization > Users.')
step(3, 'Select the user.')
step(4, 'Under Groups, select: User, Operator, Reviewer, or Silrec Admin.')
step(5, 'Save the user.')
note_box('A user can belong to multiple groups. For example, an Operator who also reviews can be assigned both Operator and Reviewer.')
screenshot('admin-user-groups.png', 'Django Admin user edit page with group checkboxes')

# ============================================================
# SCREENSHOT CHECKLIST
# ============================================================
doc.add_page_break()
h1('Screenshot Checklist')
body('Place your screenshots in the docs/screenshots/ directory. Below is every screenshot reference in this guide, grouped by section.')

checklist_sections = [
    ('1. Getting Started', [
        ('login-page.png', 'SILREC login screen'),
        ('navigation-menu.png', 'Top navigation bar'),
    ]),
    ('2. The Workflow', [
        ('workflow-diagram.png', 'Visual workflow diagram'),
        ('workflow-buttons.png', 'Workflow action buttons'),
    ]),
    ('3. New Application', [
        ('new-proposal-button.png', 'New Proposal button'),
        ('shapefile-upload.png', 'Shapefile upload section'),
        ('shapefile-validation.png', 'Shapefile attribute validation'),
        ('process-shapefile-button.png', 'Process Shapefile button'),
        ('processing-lock-icon.png', 'Lock icon during processing'),
        ('keep-revert-buttons.png', 'Keep / Revert buttons'),
        ('proposals-datatable.png', 'Proposals datatable'),
        ('proposal-detail.png', 'Proposal detail page'),
    ]),
    ('4. Entering Data', [
        ('cohort-form.png', 'Cohort edit form'),
        ('treatment-form.png', 'Treatment edit form'),
        ('prescription-form.png', 'Prescription form'),
        ('silviculturist-comments.png', 'Silviculturist comments'),
        ('survey-documents-upload.png', 'Survey document upload'),
        ('admin-form-validation-rules.png', 'Form validation rules admin'),
    ]),
    ('5. Treatments', [
        ('treatments-datatable.png', 'Treatments datatable'),
        ('treatment-detail.png', 'Treatment detail page'),
        ('treatment-extras-form.png', 'Treatment extras form'),
    ]),
    ('6. Map', [
        ('map-view.png', 'Map view with filters'),
        ('map-feature-popup.png', 'Feature popup'),
        ('merge-polygon-tool.png', 'Merge polygon tool'),
        ('cut-polygon-tool.png', 'Cut polygon tool'),
    ]),
    ('7. Reports', [
        ('report-step1-select.png', 'Report Step 1'),
        ('report-step2-parameters.png', 'Report Step 2'),
        ('report-step3-format.png', 'Report Step 3'),
        ('report-preview-modal.png', 'Preview Data modal'),
        ('admin-sql-reports.png', 'SQL Reports admin'),
        ('admin-sql-report-edit.png', 'SQL Report edit form'),
        ('generated-reports-list.png', 'Generated reports list'),
    ]),
    ('8. Search', [
        ('search-text.png', 'Text Search page'),
        ('search-text-results.png', 'Text Search results'),
        ('search-user.png', 'User Search page'),
        ('search-reference.png', 'Reference Number Search'),
    ]),
    ('9. System Maintenance', [
        ('system-maintenance.png', 'System Maintenance page'),
        ('db-dumps-list.png', 'DB Dumps list'),
        ('admin-user-groups.png', 'Admin user groups'),
    ]),
]

for section_name, items in checklist_sections:
    h3(section_name)
    add_table(
        ['#', 'File', 'Description'],
        [[str(i+1), f'`{fn}`', desc] for i, (fn, desc) in enumerate(items)],
        col_widths=[0.4, 2.5, 3.5]
    )

# Save
output_path = '/home/ubuntu/projects/silrec/silrec_tutorial.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
