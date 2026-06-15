from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(16)
section.page_height = Inches(9)
section.top_margin = Inches(0.3)
section.bottom_margin = Inches(0.3)
section.left_margin = Inches(0.3)
section.right_margin = Inches(0.3)

HEADER_BG = '2C3E50'
WHITE = 'FFFFFF'
BLACK = '000000'
GRAY = '7F8C8D'
LIGHT_BG = 'F0F2F5'
ACTION_BG = 'E8F0FE'

STATUS_HEX = {
    'Draft': '95A5A6',
    'Processing Shapefile': '5DADEC',
    'With Operator': '2ECC71',
    'With Reviewer': 'F39C12',
    'Review Completed': '8E44AD',
}

STATUSES = [
    ('Draft', 'Draft', ['Operator can upload &\nprocess shapefile']),
    ('Processing Shapefile', 'Processing\nShapefile', ['System processes\nshapefile', 'Operator / Silrec Admin\ncan send to\nWith Operator']),
    ('With Operator', 'With\nOperator', ['Operator completes\ndata entry\n(cohorts, treatments)', 'Operator can send\nto With Reviewer', 'Operator can return\nto Draft']),
    ('With Reviewer', 'With\nReviewer', ['Reviewer reviews and\napproves/rejects', 'Reviewer can send back\nto With Operator', 'Reviewer can move to\nReview Completed']),
    ('Review Completed', 'Review\nCompleted', ['Final status -\nproposal closed', 'Only Silrec Admin\ncan reopen to\nWith Reviewer']),
]

ROLES = [
    ('User', 'View-only access. Can see proposals but cannot change status.'),
    ('Operator', 'Can upload shapefiles, enter data, send to reviewer. Core data entry role.'),
    ('Reviewer', 'Can review, approve, or return proposals to operator.'),
    ('Silrec Admin', 'Full access. Can act as any role and override status transitions.'),
]

TRANSITIONS = [
    (0, 1, 'Upload & process\nshapefile', 'Operator'),
    (1, 2, 'Processing complete\n\u2192 With Operator', 'Operator,\nSilrec Admin'),
    (2, 3, 'Send to\nReviewer', 'Operator,\nSilrec Admin'),
    (2, 0, 'Return to\nDraft', 'Operator,\nSilrec Admin'),
    (3, 4, 'Approve \u2192\nReview Completed', 'Reviewer,\nSilrec Admin'),
    (3, 2, 'Send back\nto Operator', 'Reviewer,\nSilrec Admin'),
    (4, 3, 'Reopen', 'Silrec Admin'),
]

def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, font_size=10, bold=False, color=BLACK, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

def set_cell_margins(cell, top=0, bottom=0, left=30, right=30):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:start w:w="{left}" w:type="dxa"/>'
        f'<w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>')
    tcPr.append(tcMar)

def merge_cells(table, row_start, row_end, col_start, col_end):
    cell = table.cell(row_start, col_start)
    for r in range(row_start, row_end + 1):
        for c in range(col_start, col_end + 1):
            if r != row_start or c != col_start:
                cell = cell.merge(table.cell(r, c))
    return cell

def add_vmerge_cell(table, row_start, row_end, col, text, font_size, bold, color, bg_color, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    cell = merge_cells(table, row_start, row_end, col, col)
    set_cell_shading(cell, bg_color)
    set_cell_text(cell, text, font_size, bold, color, alignment)
    cell.vertical_alignment = 1
    return cell

# ============ TITLE ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SILREC \u2014 Processing Status & Role Workflow')
run.font.size = Pt(30)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string(WHITE)
pPr = p._p.get_or_add_pPr()
shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEADER_BG}" w:val="clear"/>')
pPr.append(shd)
p.paragraph_format.space_before = Pt(30)
p.paragraph_format.space_after = Pt(4)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Proposal Lifecycle: status transitions and the roles that perform them')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)
ppr2 = p2._p.get_or_add_pPr()
shd2 = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEADER_BG}" w:val="clear"/>')
ppr2.append(shd2)
p2.paragraph_format.space_after = Pt(20)

# ============ SWIMLANE TABLE ============
num = len(STATUSES)
# Each column is a status lane with: 1 header row + actions rows
max_actions = max(len(a) for _, _, a in STATUSES)
rows_per_col = 1 + max_actions  # header + action rows

swim_table = doc.add_table(rows=rows_per_col, cols=num)
swim_table.alignment = WD_TABLE_ALIGNMENT.CENTER

avail_w = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
col_w = int(avail_w / num * 914400)

for ci in range(num):
    swim_table.columns[ci].width = Emu(col_w)

# Fill each column
for ci, (status_key, status_display, actions) in enumerate(STATUSES):
    # Header row
    cell = swim_table.cell(0, ci)
    set_cell_shading(cell, STATUS_HEX[status_key])
    set_cell_text(cell, status_display, font_size=12, bold=True, color=WHITE)
    set_cell_margins(cell, top=40, bottom=40, left=30, right=30)
    cell.vertical_alignment = 1

    # Actions (remaining rows)
    for ai, act_text in enumerate(actions):
        cell = swim_table.cell(1 + ai, ci)
        set_cell_shading(cell, ACTION_BG)
        set_cell_text(cell, act_text, font_size=8, bold=False, color=BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_margins(cell, top=25, bottom=25, left=40, right=40)
        cell.vertical_alignment = 1

    # Fill remaining empty cells in column
    for ai in range(len(actions), max_actions):
        cell = swim_table.cell(1 + ai, ci)
        set_cell_shading(cell, LIGHT_BG)
        set_cell_text(cell, '', font_size=8)
        set_cell_margins(cell, top=15, bottom=15, left=30, right=30)

doc.add_paragraph()

# ============ TRANSITIONS ARROWS AS TABLE ============
p = doc.add_paragraph()
run = p.add_run('Transitions Between Statuses')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string(BLACK)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(4)

trans_table = doc.add_table(rows=1 + len(TRANSITIONS), cols=3)
trans_table.alignment = WD_TABLE_ALIGNMENT.CENTER
trans_table.style = 'Table Grid'

for i, h in enumerate(['From \u2192 To', 'Transition', 'Permitted Roles']):
    cell = trans_table.cell(0, i)
    set_cell_shading(cell, HEADER_BG)
    set_cell_text(cell, h, font_size=10, bold=True, color=WHITE)

trans_table.columns[0].width = Emu(int(col_w * 0.7))
trans_table.columns[1].width = Emu(int(col_w * 1.0))
trans_table.columns[2].width = Emu(int(col_w * 0.8))

for idx, (from_s, to_s, label, roles) in enumerate(TRANSITIONS):
    row = trans_table.rows[idx + 1]
    from_name = STATUSES[from_s][1] if isinstance(STATUSES[from_s][1], str) else STATUSES[from_s][0]
    to_name = STATUSES[to_s][1] if isinstance(STATUSES[to_s][1], str) else STATUSES[to_s][0]
    set_cell_text(row.cells[0], f'{from_name} \u2192 {to_name}', font_size=9, bold=True, color=BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(row.cells[1], label.replace('\n', ' '), font_size=9, color=BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(row.cells[2], roles, font_size=9, bold=True, color=BLACK, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for ci in range(3):
        set_cell_margins(row.cells[ci], top=15, bottom=15, left=50, right=50)
    if idx % 2 == 0:
        for ci in range(3):
            set_cell_shading(row.cells[ci], 'F8F9FA')

doc.add_paragraph()

# ============ ROLE DEFINITIONS ============
p = doc.add_paragraph()
run = p.add_run('Role Definitions')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string(BLACK)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(4)

ROLE_HEX = {
    'User': 'BDC3C7',
    'Operator': '2ECC71',
    'Reviewer': 'F39C12',
    'Silrec Admin': 'E74C3C',
}

role_table = doc.add_table(rows=1 + len(ROLES), cols=2)
role_table.alignment = WD_TABLE_ALIGNMENT.CENTER
role_table.style = 'Table Grid'

for i, h in enumerate(['Role', 'Description']):
    cell = role_table.cell(0, i)
    set_cell_shading(cell, HEADER_BG)
    set_cell_text(cell, h, font_size=10, bold=True, color=WHITE)

role_table.columns[0].width = Emu(int(col_w * 0.8))
role_table.columns[1].width = Emu(int(col_w * 2.2))

for idx, (role_name, role_desc) in enumerate(ROLES):
    row = role_table.rows[idx + 1]
    set_cell_shading(row.cells[0], ROLE_HEX[role_name])
    set_cell_text(row.cells[0], role_name, font_size=11, bold=True, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(row.cells[1], role_desc, font_size=10, color=BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    for ci in range(2):
        set_cell_margins(row.cells[ci], top=20, bottom=20, left=60, right=60)

doc.add_paragraph()

# ============ NOTE ============
p = doc.add_paragraph()
run = p.add_run('Note: ')
run.font.size = Pt(10)
run.font.bold = True
run2 = p.add_run('A user can belong to multiple groups. For example, an Operator who also reviews can be assigned both Operator and Reviewer groups.')
run2.font.size = Pt(10)

output_path = '/home/ubuntu/projects/silrec/workflow.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
